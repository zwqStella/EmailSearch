"""Per-content-type extractors. Returns plain text or '' on unsupported / failure.

Each extractor is wrapped in `extract_attachment` which produces an
`(extracted_text, status)` tuple — never raises.
"""

from __future__ import annotations

import io
import logging
from collections.abc import Callable

from emailsearch.config import get_settings
from emailsearch.db.models import ExtractionStatus

log = logging.getLogger(__name__)


# ---------- per-format helpers --------------------------------------------


def _extract_pdf(data: bytes) -> str:
    import pymupdf

    parts: list[str] = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"# {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                parts.append(",".join(cells))
    return "\n".join(parts).strip()


def _extract_plain_text(data: bytes) -> str:
    # Try the common encodings before giving up.
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def _extract_image(data: bytes) -> str:
    if not get_settings().ocr_enabled:
        return ""
    from emailsearch.extract.ocr import ocr_image

    return ocr_image(data).strip()


# ---------- registry ------------------------------------------------------


_HANDLERS: dict[str, Callable[[bytes], str]] = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _extract_xlsx,
    "text/plain": _extract_plain_text,
    "text/markdown": _extract_plain_text,
    "text/csv": _extract_plain_text,
    "text/html": _extract_plain_text,  # crude, fine for v1
}


def _resolve_handler(content_type: str) -> Callable[[bytes], str] | None:
    if not content_type:
        return None
    ct = content_type.split(";", 1)[0].strip().lower()
    if ct in _HANDLERS:
        return _HANDLERS[ct]
    if ct.startswith("image/"):
        return _extract_image
    if ct.startswith("text/"):
        return _extract_plain_text
    return None


def extract_attachment(
    content_type: str,
    data: bytes,
) -> tuple[str, ExtractionStatus, str | None]:
    """Returns (text, status, error). Never raises."""
    handler = _resolve_handler(content_type)
    if handler is None:
        return "", "unsupported", None
    try:
        text = handler(data)
    except Exception as exc:
        log.warning("extract: %s failed: %s", content_type, exc)
        return "", "failed", str(exc)
    if not text:
        return "", "empty", None
    return text, "ok", None


def extract_image_bytes(data: bytes) -> str:
    """OCR helper used by inline_images.py (already gated by config)."""
    return _extract_image(data)
