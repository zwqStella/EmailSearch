"""Extraction tests with on-the-fly fixtures (no test data files on disk)."""

from __future__ import annotations

import io
from typing import Any
from unittest.mock import patch

import pytest
from PIL import Image, ImageDraw

from emailsearch.extract.extractors import extract_attachment
from emailsearch.extract.inline_images import augment_body_with_ocr
from emailsearch.extract.pipeline import extract_email

# ---------------- fixtures -------------------------------------------------


def _make_png(text: str = "hello", size: tuple[int, int] = (200, 200)) -> bytes:
    img = Image.new("RGB", size, color="white")
    d = ImageDraw.Draw(img)
    d.text((10, 80), text, fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_pdf(text: str = "Quarterly report: revenue grew 17%.") -> bytes:
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    out = io.BytesIO()
    doc.save(out)
    doc.close()
    return out.getvalue()


def _make_docx(text: str = "DOCX hello world.") -> bytes:
    from docx import Document

    d = Document()
    d.add_paragraph(text)
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "row-cell-A"
    table.rows[0].cells[1].text = "row-cell-B"
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _make_xlsx(values: list[list[Any]] | None = None) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    for row in values or [["alpha", 1], ["beta", 2]]:
        ws.append(row)
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ---------------- extractors: format coverage -----------------------------


def test_extract_pdf_returns_text() -> None:
    text, status, err = extract_attachment("application/pdf", _make_pdf("Hello PDF"))
    assert status == "ok"
    assert "Hello PDF" in text
    assert err is None


def test_extract_docx_returns_text() -> None:
    data = _make_docx("Word body content")
    text, status, _ = extract_attachment(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data,
    )
    assert status == "ok"
    assert "Word body content" in text
    assert "row-cell-A" in text


def test_extract_xlsx_returns_text() -> None:
    text, status, _ = extract_attachment(
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        _make_xlsx([["alpha", 1], ["beta", 2]]),
    )
    assert status == "ok"
    assert "Sheet1" in text
    assert "alpha" in text


def test_extract_plain_text_utf8() -> None:
    text, status, _ = extract_attachment("text/plain", "héllo world".encode())
    assert status == "ok"
    assert text == "héllo world"


def test_extract_unsupported_type() -> None:
    text, status, _ = extract_attachment("application/x-binary-blob", b"\x00\x01")
    assert status == "unsupported"
    assert text == ""


def test_extract_failed_returns_status() -> None:
    text, status, err = extract_attachment("application/pdf", b"not really a pdf")
    assert status == "failed"
    assert text == ""
    assert err  # has a message


# ---------------- inline image OCR splicer --------------------------------


def test_augment_body_with_ocr_splices_text(monkeypatch: pytest.MonkeyPatch) -> None:
    """When OCR is enabled and a cid: image is present, OCR text is spliced in."""

    # Stub out the actual OCR engine so the test is fast and deterministic.
    monkeypatch.setattr(
        "emailsearch.extract.extractors.extract_image_bytes",
        lambda b: "OCRED-TEXT" if b else "",
    )

    body_html = '<p>before</p><img src="cid:foo"><p>after</p>'
    inline_atts = [
        {
            "isInline": True,
            "contentId": "foo",
            "_bytes": b"fake-png-bytes",
        }
    ]
    text, ocr_used = augment_body_with_ocr(body_html, inline_atts)
    assert ocr_used is True
    assert "before" in text
    assert "after" in text
    assert "OCRED-TEXT" in text


def test_augment_body_no_cid_returns_plain_text(monkeypatch: pytest.MonkeyPatch) -> None:
    """No cid: refs → no OCR, just html2text."""
    body_html = "<p>just text</p>"
    text, ocr_used = augment_body_with_ocr(body_html, [])
    assert ocr_used is False
    assert "just text" in text


def test_augment_body_skips_when_ocr_disabled() -> None:
    from emailsearch.config import Settings

    # patch settings via the module-level cache
    with patch("emailsearch.extract.inline_images.get_settings") as mock:
        mock.return_value = Settings(
            data_dir=__import__("pathlib").Path("."),
            ocr_enabled=False,
        )
        text, ocr_used = augment_body_with_ocr(
            '<p>x</p><img src="cid:foo">',
            [{"isInline": True, "contentId": "foo", "_bytes": b"x"}],
        )
    assert ocr_used is False
    assert "x" in text


# ---------------- end-to-end pipeline -------------------------------------


def test_extract_email_assembles_record(monkeypatch: pytest.MonkeyPatch) -> None:
    # Stub OCR so the inline image path is exercised without invoking the real engine.
    monkeypatch.setattr(
        "emailsearch.extract.extractors.extract_image_bytes",
        lambda b: "INLINE-OCR",
    )

    from emailsearch.outlook.raw import RawAttachment, RawMessage

    pdf_bytes = _make_pdf("PDF says: budget approved.")
    inline_png = _make_png("inline")
    raw = RawMessage(
        id="msg-1",
        subject="Hello",
        from_address="alice@example.com",
        from_name="Alice",
        to=[("bob@example.com", "Bob")],
        cc=[],
        received_at=1736942400,  # 2025-01-15T12:00:00Z
        sent_at=1736942340,
        body_html='<p>See pic:</p><img src="cid:img1"><p>End.</p>',
        body_preview="See pic",
        conversation_id="conv-1",
        folder_id="folder-x",
        web_link="outlook:msg-1",
        has_attachments=True,
        attachments=[
            RawAttachment(
                att_id="att-pdf",
                name="report.pdf",
                content_type="application/pdf",
                size=len(pdf_bytes),
                is_inline=False,
                content_bytes=pdf_bytes,
            ),
            RawAttachment(
                att_id="att-img",
                name="logo.png",
                content_type="image/png",
                size=len(inline_png),
                is_inline=True,
                content_id="img1",
                content_bytes=inline_png,
            ),
            RawAttachment(
                att_id="att-big",
                name="huge.bin",
                content_type="application/octet-stream",
                size=9_999_999,
                is_inline=False,
                content_bytes=None,
                skipped_reason="too_large:9999999>26214400",
            ),
        ],
    )

    email = extract_email(raw)

    assert email.id == "msg-1"
    assert email.subject == "Hello"
    assert email.from_address == "alice@example.com"
    assert email.to_addresses[0].address == "bob@example.com"
    assert email.received_at > 0
    assert email.body_ocr_used is True
    assert "INLINE-OCR" in email.body_text
    assert "See pic" in email.body_text
    assert email.has_attachments is True
    by_name = {a.name: a for a in email.attachments}
    assert by_name["report.pdf"].status == "ok"
    assert "budget approved" in by_name["report.pdf"].extracted_text
    assert by_name["logo.png"].status in ("ok", "empty")  # tiny rendered text → may be empty
    assert by_name["huge.bin"].status == "skipped_too_large"
    # searchable_text concatenates body + each attachment's extracted_text
    assert "budget approved" in email.searchable_text
    assert "INLINE-OCR" in email.searchable_text
